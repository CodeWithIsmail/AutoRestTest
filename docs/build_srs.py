#!/usr/bin/env python
"""Assemble the complete AutoRestTest SRS report into a single .docx.

Reads the ordered section markdown files in docs/srs-content/, renders them with
a small Markdown subset (headings, paragraphs, bullet lists, block-quotes, GFM
tables, inline **bold**/*italic*/`code`, and `@fig file | caption` figure
directives), and emits a Word document with a title page, an auto-updating table
of contents, page numbers, and captioned figures scaled to the page.

Run:  python docs/build_srs.py
Deps: python-docx, Pillow   (both already present)
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parent
CONTENT = ROOT / "srs-content"
PNG = ROOT / "srs-diagrams" / "png"
OUT = ROOT.parent / "resources" / "SRS" / "AutoRestTest SRS.docx"

SECTION_FILES = [
    "01-introduction.md",
    "02-qfd.md",
    "03-usage-scenario.md",
    "04-use-case-diagrams.md",
    "05-activity-diagrams.md",
    "06-data-modeling.md",
]

NAVY = RGBColor(0x1F, 0x39, 0x64)
GREY = RGBColor(0x44, 0x44, 0x44)
MAX_W_IN = 6.3   # content width at 1" margins on Letter
MAX_H_IN = 8.4   # leave room for the caption on one page

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*)")

_fig_no = [0]


# ---------------------------------------------------------------------------
# low-level helpers
# ---------------------------------------------------------------------------
def add_runs(paragraph, text: str):
    """Add `text` to a paragraph, honouring **bold**, *italic*, `code`."""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:  # *italic*
            paragraph.add_run(tok[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _fld(runtext: str, instr: str):
    """Build the OOXML run children for a Word field (begin/instr/sep/end)."""
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = runtext
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    return [begin, it, sep, t, end]


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    for el in _fld("Right-click here and choose “Update Field” to build the table of contents.",
                   r'TOC \o "1-3" \h \z \u'):
        run._r.append(el)


def add_page_numbers(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True  # no number on title page
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for el in _fld("1", "PAGE"):
        run._r.append(el)


def add_figure(doc, path: Path, caption: str):
    if not path.exists():
        raise FileNotFoundError(path)
    with Image.open(path) as im:
        w, h = im.size
    aspect = h / w
    width = MAX_W_IN
    if width * aspect > MAX_H_IN:
        width = MAX_H_IN / aspect

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(str(path), width=Inches(width))

    _fig_no[0] += 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(f"Figure {_fig_no[0]}: {caption}")
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY


def _set_repeat_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true")
    trPr.append(th)


def add_table(doc, block: list[str]):
    def cells(line):
        return [c.strip() for c in line.strip().strip("|").split("|")]

    header = cells(block[0])
    rows = [cells(l) for l in block[2:]]  # block[1] is the |---| separator
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
    _set_repeat_header(table.rows[0])

    for r in rows:
        tcells = table.add_row().cells
        for i in range(len(header)):
            val = r[i] if i < len(r) else ""
            para = tcells[i].paragraphs[0]
            para.text = ""
            add_runs(para, val)
            for run in para.runs:
                run.font.size = Pt(9.5)


def add_note(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    # Render inline formatting, then italicise everything that isn't bold.
    add_runs(p, text)
    for run in p.runs:
        if not run.bold:
            run.italic = True


# ---------------------------------------------------------------------------
# markdown block parser
# ---------------------------------------------------------------------------
SPECIAL = ("#", ">", "|", "@fig ", "- ")


def is_special(s: str) -> bool:
    return (not s) or s == "---" or s.startswith(SPECIAL)


def process_markdown(doc, text: str):
    lines = text.split("\n")
    i, n = 0, len(lines)
    while i < n:
        s = lines[i].strip()
        if not s or s == "---":
            i += 1
            continue

        if s.startswith("@fig "):
            fname, caption = s[5:].split("|", 1)
            add_figure(doc, PNG / fname.strip(), caption.strip())
            i += 1
            continue

        if s.startswith("#"):
            m = re.match(r"(#+)\s+(.*)", s)
            level = min(len(m.group(1)), 4)
            h = doc.add_heading("", level=level)
            add_runs(h, m.group(2))
            i += 1
            continue

        if s.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip()[1:].strip())
                i += 1
            add_note(doc, " ".join(buf))
            continue

        if s.startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            add_table(doc, block)
            continue

        if s.startswith("- "):
            while i < n and lines[i].strip().startswith("- "):
                text = lines[i].strip()[2:]
                i += 1
                # absorb hard-wrapped continuation lines of this bullet
                while i < n:
                    cont = lines[i].strip()
                    if is_special(cont):
                        break
                    text += " " + cont
                    i += 1
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, text)
            continue

        # normal paragraph: gather wrapped lines until a blank/special line
        buf = []
        while i < n and not is_special(lines[i].strip()):
            buf.append(lines[i].strip())
            i += 1
        add_runs(doc.add_paragraph(), " ".join(buf))


# ---------------------------------------------------------------------------
# document scaffold
# ---------------------------------------------------------------------------
def style_base(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for lvl, size in ((1, 18), (2, 14), (3, 12), (4, 11)):
        st = doc.styles[f"Heading {lvl}"]
        st.font.color.rgb = NAVY
        st.font.size = Pt(size)


def title_page(doc):
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AutoRestTest"); r.bold = True; r.font.size = Pt(40); r.font.color.rgb = NAVY

    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("An AI-Powered Collaborative Platform for Automated REST API Testing")
    r.italic = True; r.font.size = Pt(13); r.font.color.rgb = GREY

    for _ in range(2):
        doc.add_paragraph()
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Software Requirement Specification\n& Analysis Report")
    r.bold = True; r.font.size = Pt(16)

    for _ in range(6):
        doc.add_paragraph()
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_run("July 2026").font.size = Pt(12)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def toc_page(doc):
    h = doc.add_paragraph()
    r = h.add_run("Table of Contents"); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
    add_toc(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def main():
    doc = Document()
    style_base(doc)
    add_page_numbers(doc)
    title_page(doc)
    toc_page(doc)

    for idx, fname in enumerate(SECTION_FILES):
        text = (CONTENT / fname).read_text(encoding="utf-8")
        process_markdown(doc, text)
        if idx != len(SECTION_FILES) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"wrote {OUT}  ({_fig_no[0]} figures)")


if __name__ == "__main__":
    main()
